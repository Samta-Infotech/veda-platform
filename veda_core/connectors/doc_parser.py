# =============================================================================
# connectors/doc_parser.py
# VEDA — Layout-aware document parsing (Cross-source plan, Phase 3)
#
# Replaces flat text extraction with STRUCTURE-aware parsing so the narrative lane
# carries a doc → section → chunk hierarchy instead of a bag of disconnected
# chunks. Emits a common ParsedDoc{metadata, sections[{path, level, text, tables}]}:
#
#   PDF  → pymupdf4llm (markdown with heading hierarchy + table detection; pure
#          local, no egress). Headings (#, ##, …) define sections; markdown pipe
#          tables are captured per section and routed to the tabular lane (P2.3).
#   DOCX → python-docx walking Heading styles + Table objects.
#
# Structure-aware chunking (chunk_sections) never crosses a section boundary and
# prepends the heading breadcrumb to each chunk's embedded text
# ("Contracts > Termination > Notice periods:\n<text>") — measurably better chunk
# retrieval precision at zero cost. Tables with ≥ DOC_TABLE_MIN_ROWS rows become
# derived tables (P2.3); smaller/ragged tables stay as chunk text.
#
# Optional deps: pymupdf4llm (PDF), python-docx (DOCX). Absent → the caller falls
# back to the flat connectors.document extractors.
# =============================================================================

from __future__ import annotations

import os
import re
from dataclasses import dataclass, field
from typing import List, Optional

from config import DOC_CHUNK_SIZE, DOC_CHUNK_OVERLAP, DOC_TABLE_MIN_ROWS

try:
    import pymupdf4llm as _pymupdf4llm
    _PYMUPDF4LLM_AVAILABLE = True
except ImportError:
    _PYMUPDF4LLM_AVAILABLE = False

try:
    from docx import Document as _DocxDocument
    _DOCX_AVAILABLE = True
except ImportError:
    _DOCX_AVAILABLE = False


@dataclass
class DocTable:
    """A table detected inside a document. ``rows`` are lists of cell strings
    (header first). Routed to the tabular lane when len(rows)-1 ≥ DOC_TABLE_MIN_ROWS."""
    header: List[str]
    rows: List[List[str]]
    section_path: str = ""

    @property
    def n_rows(self) -> int:
        return len(self.rows)

    def is_derived_table(self) -> bool:
        return bool(self.header) and self.n_rows >= DOC_TABLE_MIN_ROWS


@dataclass
class Section:
    path: str          # heading breadcrumb, e.g. "Contracts > Termination > Notice"
    level: int
    text: str
    tables: List[DocTable] = field(default_factory=list)


@dataclass
class ParsedDoc:
    metadata: dict
    sections: List[Section] = field(default_factory=list)

    def derived_tables(self) -> List[DocTable]:
        return [t for s in self.sections for t in s.tables if t.is_derived_table()]


def parser_available(ext: str) -> bool:
    ext = ext.lower().lstrip(".")
    if ext == "pdf":
        return _PYMUPDF4LLM_AVAILABLE
    if ext == "docx":
        return _DOCX_AVAILABLE
    return False


# --------------------------------------------------------------------- markdown
_MD_TABLE_ROW = re.compile(r"^\s*\|(.+)\|\s*$")
_MD_TABLE_SEP = re.compile(r"^\s*\|?[\s:-]+\|[\s:|-]*$")


def _split_md_row(line: str) -> List[str]:
    inner = line.strip().strip("|")
    return [c.strip() for c in inner.split("|")]


def parse_markdown(md: str, metadata: dict) -> ParsedDoc:
    """Parse a markdown string (pymupdf4llm output) into a ParsedDoc: ATX headings
    build the section breadcrumb stack; consecutive pipe-table lines become DocTables
    attached to the current section."""
    sections: List[Section] = []
    stack: List[tuple] = []          # (level, title)
    cur = Section(path="(root)", level=0, text="")
    sections.append(cur)
    lines = md.splitlines()
    i = 0
    buf: List[str] = []

    def _flush():
        cur.text = "\n".join(buf).strip()

    while i < len(lines):
        line = lines[i]
        h = re.match(r"^(#{1,6})\s+(.*)$", line)
        if h:
            _flush()
            level = len(h.group(1))
            title = h.group(2).strip()
            while stack and stack[-1][0] >= level:
                stack.pop()
            stack.append((level, title))
            path = " > ".join(t for _, t in stack)
            cur = Section(path=path, level=level, text="")
            sections.append(cur)
            buf = []
            i += 1
            continue
        # markdown table: header row, separator row, then body rows
        if _MD_TABLE_ROW.match(line) and i + 1 < len(lines) and _MD_TABLE_SEP.match(lines[i + 1]):
            header = _split_md_row(line)
            body: List[List[str]] = []
            j = i + 2
            while j < len(lines) and _MD_TABLE_ROW.match(lines[j]) and not _MD_TABLE_SEP.match(lines[j]):
                body.append(_split_md_row(lines[j]))
                j += 1
            cur.tables.append(DocTable(header=header, rows=body, section_path=cur.path))
            i = j
            continue
        buf.append(line)
        i += 1
    _flush()
    return ParsedDoc(metadata=metadata, sections=[s for s in sections if s.text or s.tables])


# ------------------------------------------------------------------------- PDF
def parse_pdf(path: str) -> Optional[ParsedDoc]:
    if not _PYMUPDF4LLM_AVAILABLE:
        return None
    try:
        md = _pymupdf4llm.to_markdown(path)
    except Exception:
        return None
    meta = {"name": os.path.basename(path), "path": path, "mime": "application/pdf",
            "author": _pdf_author(path)}
    return parse_markdown(md, meta)


def _pdf_author(path: str) -> str:
    try:
        import pymupdf  # fitz, pulled in by pymupdf4llm
        with pymupdf.open(path) as d:
            return (d.metadata or {}).get("author", "") or ""
    except Exception:
        return ""


# ------------------------------------------------------------------------ DOCX
def parse_docx(path: str) -> Optional[ParsedDoc]:
    if not _DOCX_AVAILABLE:
        return None
    try:
        doc = _DocxDocument(path)
    except Exception:
        return None
    sections: List[Section] = []
    stack: List[tuple] = []
    cur = Section(path="(root)", level=0, text="")
    sections.append(cur)
    buf: List[str] = []

    def _flush():
        cur.text = "\n".join(buf).strip()

    for p in doc.paragraphs:
        style = (p.style.name if p.style else "") or ""
        m = re.match(r"Heading\s+(\d+)", style)
        if m and p.text.strip():
            _flush()
            level = int(m.group(1))
            while stack and stack[-1][0] >= level:
                stack.pop()
            stack.append((level, p.text.strip()))
            cur = Section(path=" > ".join(t for _, t in stack), level=level, text="")
            sections.append(cur)
            buf = []
        elif p.text.strip():
            buf.append(p.text)
    _flush()
    # tables → attach to the last section (python-docx tables are document-ordered)
    for t in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        if rows:
            (sections[-1] if sections else cur).tables.append(
                DocTable(header=rows[0], rows=rows[1:], section_path=sections[-1].path))
    meta = {"name": os.path.basename(path), "path": path,
            "mime": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "author": (doc.core_properties.author or "") if doc.core_properties else ""}
    return ParsedDoc(metadata=meta, sections=[s for s in sections if s.text or s.tables])


def parse_document(path: str) -> Optional[ParsedDoc]:
    ext = os.path.splitext(path)[1].lower().lstrip(".")
    if ext == "pdf":
        return parse_pdf(path)
    if ext == "docx":
        return parse_docx(path)
    return None


# ------------------------------------------------------------- structure chunking
def _word_chunks(text: str, size: int, overlap: int) -> List[str]:
    words = text.split()
    if not words:
        return []
    step = max(1, size - overlap)
    out = []
    for start in range(0, len(words), step):
        seg = " ".join(words[start:start + size])
        if seg:
            out.append(seg)
        if start + size >= len(words):
            break
    return out


def chunk_sections(parsed: ParsedDoc, chunk_size: int = DOC_CHUNK_SIZE,
                   chunk_overlap: int = DOC_CHUNK_OVERLAP) -> List[dict]:
    """Structure-aware chunks that never cross a section boundary and carry the
    heading path both as embedded-text prefix and as metadata. Returns dicts:
    {text, embed_text, section_path, section_level, chunk_index}."""
    out: List[dict] = []
    idx = 0
    for sec in parsed.sections:
        if not sec.text:
            continue
        prefix = "" if sec.path in ("", "(root)") else f"{sec.path}:\n"
        for seg in _word_chunks(sec.text, chunk_size, chunk_overlap):
            out.append({"text": seg, "embed_text": prefix + seg,
                        "section_path": "" if sec.path == "(root)" else sec.path,
                        "section_level": sec.level, "chunk_index": idx})
            idx += 1
    return out
